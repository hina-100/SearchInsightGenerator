# -*- coding: utf-8 -*-
"""
Created on Sat Apr  6 00:34:56 2024

@author: LENOVO
"""
# Outline
# Creating a rest API to fetch top 10 search info from the second layer
# Fetching Data from Json and using NLP to summarize the info
# Creating an API to push on the result to UI

import json
import nltk

from nltk.tokenize import sent_tokenize
from heapq import nlargest
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import networkx as nx

# Load the local json file

with open('data.json','r') as file:
    data = json.load(file)
    
    
# Loading data from Mongo DB
from pymongo import MongoClient

# Connect to MongoDB Cloud server
client = MongoClient("mongodb+srv://hinaagrawal100:p6X6OypCUMIhi1dA@cluster0.0mped2d.mongodb.net/?retryWrites=true&w=majority&appName=Cluster0")

# Select database
db = client["Search"]

# Select collection
collection = db["SearchHistory"]

New_Collection=db["Advance_search_History"]


# Print documents
# for doc in documents:
#     print(doc)
    
# List to store documents
document_list = []

# Find documents
documents = collection.find()

for doc in documents:
    document_list.append(doc)    
    
# Extract snippets from the organic section

snippets = [result.get('snippet','') for result in data.get('organic',[])]   

# Combine all snippets into a single text
data = document_list[2]


# Extracting relevant information
items = data["SearchResult"]["items"]
search_results = []

for item in items:
    title = item["title"]
    link = item["link"]
    snippet = item["snippet"]
    search_results.append({"title": title, "link": link, "snippet": snippet})


## preprocess the input text
#preprocessed_text = text.strip().replace('\n','')

text = ''.join(snippets)



print(text)

# Tokenize the text into sentences

sentences = sent_tokenize(text)

# Initialize NLTK's punkt tokenizer (required for sentence tokenization)

nltk.download('punkt')

# Calculate sentence scores based on the number of words in each sentence

sentence_scores = {sentence: len(sentence.split()) for sentence in sentences}

# Get the top 5 most relevant sentences based on score 

"""Approach 1"""

summary_sentences= nlargest(5, sentence_scores, key= sentence_scores.get)

"""Approach2- TF-IDF"""
# Initialize the TF-IDF vectorizer

vectorizer = TfidfVectorizer()
# Compute TF-IDF scores for each sentence

tfidf_matrix = vectorizer.fit_transform([text])
feature_names =  vectorizer.get_feature_names_out()

# Calculate sentence scores based on TF-IDF score of words

sentence_scores ={}

for i, sentence in enumerate(sentences):
    score= sum(tfidf_matrix[0,vectorizer.vocabulary_[word]] for word in sentence.split() if word in vectorizer.vocabulary_)

# Get the top 5 most relevant sentences based on score
summary_sentences = nlargest(5, sentence_scores, key=sentence_scores.get)

"""Approach3- Text Rank"""


# Compute similarity matrix between sentences
similarity_matrix = cosine_similarity(tfidf_matrix, tfidf_matrix)

# Apply TextRank algorithm
graph = nx.from_numpy_array(similarity_matrix)
scores = nx.pagerank(graph)

# Filter out None values from scores
filtered_scores = {k: v for k, v in scores.items() if v is not None}

# Sort sentences by TextRank scores
summary_sentences = nlargest(5, filtered_scores, key=filtered_scores.get)




"""Approach4- BERT Embeddings"""


from sentence_transformers import SentenceTransformer

# Load pre-trained BERT model
model = SentenceTransformer('bert-base-nli-mean-tokens')

# Generate sentence embeddings
sentence_embeddings = model.encode(sentences)

# Compute cosine similarity matrix
similarity_matrix = cosine_similarity(sentence_embeddings)

# Calculate sentence scores based on similarity to other sentences
scores = similarity_matrix.sum(axis=1)

# Get the top 5 most relevant sentences based on score
summary_sentences = nlargest(5, zip(sentences, scores), key=lambda x: x[1])
summary_sentences = [sentence for sentence, _ in summary_sentences]



"""Approach5- Sentence position and length"""

# Calculate sentence scores based on position and length
sentence_scores = {}
for i, sentence in enumerate(sentences):
    score = (1 / (i + 1)) * (1 / len(sentence.split()))
    sentence_scores[sentence] = score

# Get the top 5 most relevant sentences based on score
summary_sentences = nlargest(5, sentence_scores, key=sentence_scores.get)


"""Approach6- NER"""
import spacy

# Load SpaCy model for NER
nlp = spacy.load("en_core_web_sm")

# Identify named entities in each sentence
entity_scores = {}
for sentence in sentences:
    doc = nlp(sentence)
    entities = [ent.text for ent in doc.ents]
    score = len(entities)
    entity_scores[sentence] = score

# Get the top 5 most relevant sentences based on named entities
summary_sentences = nlargest(5, entity_scores, key=entity_scores.get)




 
#Print the summary

print("Summary:")

for sentence in summary_sentences:
    print(sentence)

summary_sentences1= summary_sentences

summary_sentences2= summary_sentences

summary_sentences_text_rank= summary_sentences
#Got no result here

summary_sentences_BERT= summary_sentences
#This took time

summary_sentences_pos_l= summary_sentences

summary_sentences_NER= summary_sentences

# Writing output to files

from docx import Document

# Create a new Document object
doc = Document()

# Add text to the document
# text = "This is some text that we want to write to a Word document."
# doc.add_paragraph(text)

# Function to add heading and sentences to the document
def add_heading_and_sentences(heading, sentences):
    doc.add_heading(heading, level=1)
    for sentence in sentences:
        doc.add_paragraph(sentence)
        
# Add headings and sentences to the document
add_heading_and_sentences("Summary Sentences 1", summary_sentences1)
#add_heading_and_sentences("Summary Sentences 2", summary_sentences2)
add_heading_and_sentences("TextRank Summary Sentences", summary_sentences_text_rank)
add_heading_and_sentences("BERT Summary Sentences", summary_sentences_BERT)
add_heading_and_sentences("Position and Length Summary Sentences", summary_sentences_pos_l)
add_heading_and_sentences("NER Summary Sentences", summary_sentences_NER)        


# Save the document
doc.save("output.docx")


###### Abstractive Summarization ##############

"""T5 (Text-to- text transfer transformer)"""

import torch
from transformers import T5Tokenizer, T5ForConditionalGeneration, T5Config 
import sentencepiece


# initialize the pretrained model
model = T5ForConditionalGeneration.from_pretrained('t5-small')
tokenizer = T5Tokenizer.from_pretrained('t5-small')
device = torch.device('cpu')

## preprocess the input text
preprocessed_text = text.strip().replace('\n','')
t5_input_text = 'summarize: ' + preprocessed_text

# There is limit on input size of 512 words. In case there are more words, it can be splitted into smaller chunks
# and there can summary created of 50 -50 words from each.

# Other alternative is to use tr-large/medium model and use GPU instead of CPU. The model 
# will be more efficient but there will be increased processing time

len(t5_input_text.split()) 

tokenized_text = tokenizer.encode(t5_input_text, return_tensors='pt', max_length=512).to(device)

summary_ids = model.generate(tokenized_text, min_length=30, max_length=120)
summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)

summary

"""BART (Bidirectional and autoregressive transformers)"""

from transformers import BartTokenizer, BartForConditionalGeneration

# Load pre-trained BART tokenizer and model
tokenizer = BartTokenizer.from_pretrained('facebook/bart-large-cnn')
model = BartForConditionalGeneration.from_pretrained('facebook/bart-large-cnn')

# Tokenize the input text
inputs = tokenizer([text], max_length=1024, return_tensors='pt', truncation=True)

# Generate summary
summary_ids = model.generate(inputs['input_ids'], num_beams=4, min_length=30, max_length=150, early_stopping=True)

# Decode the summary tokens back into text
summary_text = tokenizer.decode(summary_ids[0], skip_special_tokens=True)

# Print the generated summary
print(summary_text)

# Writing Result to Mongo DB

# Insert a document
document = {"Summary": summary_text}
New_Collection.insert_one(document)

# Close connection
client.close()


